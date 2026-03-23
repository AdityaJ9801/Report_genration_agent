"""
python-docx based DOCX exporter.
Generates Word documents with heading styles, embedded charts,
formatted data tables, and page breaks between sections.
"""

import base64
import io
import logging
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from app.models import NarrativeSections, BrandingConfig

logger = logging.getLogger(__name__)


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex color string to RGBColor."""
    hex_color = hex_color.lstrip("#")
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


def _add_styled_heading(doc: Document, text: str, level: int, color: Optional[RGBColor] = None):
    """Add a heading with optional color styling."""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color


def _add_formatted_table(doc: Document, data: list[dict], accent_color: Optional[RGBColor] = None):
    """Add a formatted table from a list of dictionaries."""
    if not data:
        return

    headers = list(data[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header).replace("_", " ").title()
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_data in data[:50]:  # Cap at 50 rows
        row = table.add_row()
        for i, header in enumerate(headers):
            cell = row.cells[i]
            cell.text = str(row_data.get(header, ""))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def _add_chart_image(doc: Document, chart: dict):
    """Embed a chart image from base64 data."""
    image_b64 = chart.get("image_base64", chart.get("png_base64", ""))
    if not image_b64:
        return

    try:
        image_bytes = base64.b64decode(image_b64)
        image_stream = io.BytesIO(image_bytes)
        doc.add_picture(image_stream, width=Inches(6.0))

        # Add caption
        title = chart.get("title", "Chart")
        caption = doc.add_paragraph(title)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.style = doc.styles["Caption"] if "Caption" in doc.styles else None
    except Exception as e:
        logger.warning(f"Failed to embed chart image: {e}")
        doc.add_paragraph(f"[Chart: {chart.get('title', 'Untitled')} — image unavailable]")


def export_docx(
    narratives: NarrativeSections,
    charts: Optional[list[dict]] = None,
    branding: Optional[BrandingConfig] = None,
    report_style: str = "detailed",
    user_query: str = "",
    include_charts: bool = True,
    sql_results: Optional[list[dict]] = None,
) -> bytes:
    """
    Generate a DOCX report.

    Returns:
        DOCX file as bytes
    """
    doc = Document()
    brand = branding or BrandingConfig()
    accent = _hex_to_rgb(brand.primary_color)

    # --- Title Page ---
    if brand.logo_base64:
        try:
            logo_bytes = base64.b64decode(brand.logo_base64)
            logo_stream = io.BytesIO(logo_bytes)
            doc.add_picture(logo_stream, width=Inches(2.0))
        except Exception:
            pass

    title = doc.add_heading("Analysis Report", level=0)
    for run in title.runs:
        run.font.color.rgb = accent

    if brand.company_name:
        subtitle = doc.add_paragraph(brand.company_name)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = accent

    if user_query:
        doc.add_paragraph(f"Query: {user_query}").italic = True

    doc.add_page_break()

    # --- Executive Summary ---
    if narratives.executive_summary:
        _add_styled_heading(doc, "Executive Summary", level=1, color=accent)
        doc.add_paragraph(narratives.executive_summary)
        doc.add_page_break()

    # --- Data Overview ---
    if narratives.data_overview:
        _add_styled_heading(doc, "Data Overview", level=1, color=accent)
        doc.add_paragraph(narratives.data_overview)
        doc.add_page_break()

    # --- SQL Findings ---
    if narratives.sql_findings:
        _add_styled_heading(doc, "SQL Analysis Findings", level=1, color=accent)
        doc.add_paragraph(narratives.sql_findings)

        # Embed SQL results table if available
        if sql_results:
            _add_styled_heading(doc, "Query Results", level=2, color=accent)
            _add_formatted_table(doc, sql_results, accent)

        doc.add_page_break()

    # --- ML Insights ---
    if narratives.ml_insights:
        _add_styled_heading(doc, "Machine Learning Insights", level=1, color=accent)
        doc.add_paragraph(narratives.ml_insights)
        doc.add_page_break()

    # --- NLP Section ---
    if narratives.nlp_section:
        _add_styled_heading(doc, "NLP & Text Analysis", level=1, color=accent)
        doc.add_paragraph(narratives.nlp_section)
        doc.add_page_break()

    # --- Charts ---
    if include_charts and charts:
        _add_styled_heading(doc, "Visualizations", level=1, color=accent)
        for chart in charts:
            _add_chart_image(doc, chart)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.read()
    logger.info(f"Generated DOCX report ({len(docx_bytes)} bytes)")
    return docx_bytes
